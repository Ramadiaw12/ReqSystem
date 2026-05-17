"""
utils/docx_exporter.py — Génération du CDC en format Word (.docx)

Produit un document professionnel avec :
  - Page de garde colorée
  - En-tête et pied de page sur chaque page
  - Tableau des informations client
  - Toutes les sections du CDC
  - Tableau des exigences avec priorités colorées
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime


# ── Couleurs ──────────────────────────────────────────────────────────────────
BLUE       = RGBColor(0x1B, 0x4F, 0x8A)
LIGHT_BLUE = RGBColor(0xD5, 0xE8, 0xF5)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
DARK_GRAY  = RGBColor(0x2C, 0x2C, 0x2C)
MED_GRAY   = RGBColor(0x55, 0x55, 0x55)
RED        = RGBColor(0xC5, 0x30, 0x30)
AMBER      = RGBColor(0xB7, 0x79, 0x1F)
GREEN      = RGBColor(0x27, 0x67, 0x49)
LIGHT_GRAY = RGBColor(0xF5, 0xF5, 0xF5)

OUTPUT_DIR = "/tmp/systemreq_exports"


def set_cell_bg(cell, hex_color: str):
    """Applique une couleur de fond à une cellule."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def add_paragraph(doc, text, bold=False, size=11, color=None, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.bold      = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold            = True
    run.font.size       = Pt(14 if level == 1 else 12)
    run.font.color.rgb  = BLUE
    # Ligne sous le titre niveau 1
    if level == 1:
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"),   "single")
        bottom.set(qn("w:sz"),    "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "1B4F8A")
        pBdr.append(bottom)
        pPr.append(pBdr)
    return p


async def export_to_docx(cdc: dict) -> str:
    """
    Génère le fichier .docx et retourne son chemin.

    Args:
        cdc : dict complet issu du GeneratorAgent

    Returns:
        str : chemin absolu vers le fichier généré
    """
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    session_id   = cdc.get("session_id", "unknown")
    client_name  = cdc.get("client_name", "Client")
    project_name = cdc.get("project_name", "Projet web")
    services     = ", ".join(cdc.get("services", []))
    budget       = cdc.get("budget", "N/A")
    sections     = cdc.get("sections", [])
    requirements = cdc.get("requirements", [])
    date_str     = datetime.now().strftime("%d/%m/%Y")
    ref          = f"CDC-{session_id[:6].upper()}"

    doc = Document()

    # ── Marges ────────────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Page de garde ─────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(60)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("EPS SARL")
    run.font.size      = Pt(13)
    run.font.color.rgb = BLUE
    run.bold = True

    add_paragraph(doc, "CAHIER DES CHARGES PRÉLIMINAIRE",
                  bold=True, size=24, color=BLUE,
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  space_before=10, space_after=10)

    add_paragraph(doc, "─" * 50, size=10, color=LIGHT_BLUE,
                  align=WD_ALIGN_PARAGRAPH.CENTER)

    add_paragraph(doc, f"Client : {client_name}",
                  size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10)

    add_paragraph(doc, f"Services : {services}",
                  size=11, color=MED_GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)

    add_paragraph(doc, f"Budget : {budget} EUR  ·  {date_str}",
                  size=11, color=MED_GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)

    add_paragraph(doc, f"Réf. {ref}",
                  size=10, color=MED_GRAY, align=WD_ALIGN_PARAGRAPH.CENTER,
                  space_before=10)

    doc.add_page_break()

    # ── Section 1 : Informations client ───────────────────────────────────────
    add_heading(doc, "1. Informations du client")

    info_rows = [
        ("Nom & Prénom",    client_name),
        ("Email",           cdc.get("email", "—")),
        ("Téléphone",       cdc.get("telephone", "—")),
        ("Ville / Pays",    f"{cdc.get('ville', '')} — {cdc.get('pays', '')}"),
        ("Budget",          f"{budget} EUR"),
        ("Services",        services),
        ("Référence",       ref),
        ("Date",            date_str),
    ]

    table = doc.add_table(rows=len(info_rows) + 1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # En-tête
    hdr = table.rows[0].cells
    hdr[0].text = "Champ"
    hdr[1].text = "Valeur"
    for cell in hdr:
        set_cell_bg(cell, "1B4F8A")
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = WHITE
            run.bold = True
            run.font.size = Pt(10)

    # Données
    for i, (label, value) in enumerate(info_rows):
        row = table.rows[i + 1].cells
        row[0].text = label
        row[1].text = str(value)
        set_cell_bg(row[0], "D5E8F5")
        for run in row[0].paragraphs[0].runs:
            run.font.color.rgb = BLUE
            run.bold = True
            run.font.size = Pt(9)
        for run in row[1].paragraphs[0].runs:
            run.font.size = Pt(9)

    doc.add_paragraph()

    # ── Sections du CDC ───────────────────────────────────────────────────────
    for section in sections:
        add_heading(doc, section.get("title", ""), level=1)
        content = section.get("content", "")
        for line in content.split("\n"):
            if line.strip():
                add_paragraph(doc, line.strip(), size=10, color=DARK_GRAY,
                              space_before=2, space_after=4)

    doc.add_page_break()

    # ── Tableau des exigences ─────────────────────────────────────────────────
    add_heading(doc, "Tableau des exigences")

    if requirements:
        req_table = doc.add_table(rows=len(requirements) + 1, cols=5)
        req_table.style   = "Table Grid"
        req_table.alignment = WD_TABLE_ALIGNMENT.LEFT

        headers = ["ID", "Description", "Catégorie", "Priorité", "MoSCoW"]
        for i, h in enumerate(headers):
            cell = req_table.rows[0].cells[i]
            cell.text = h
            set_cell_bg(cell, "1B4F8A")
            for run in cell.paragraphs[0].runs:
                run.font.color.rgb = WHITE
                run.bold = True
                run.font.size = Pt(9)

        for i, req in enumerate(requirements):
            row = req_table.rows[i + 1].cells
            row[0].text = req.get("id", "")
            row[1].text = req.get("description", "")
            row[2].text = req.get("category", "")
            row[3].text = req.get("priority", "")
            row[4].text = req.get("moscow", "")

            # Couleur de priorité
            priority = req.get("priority", "")
            color_map = {"Haute": "FFF0F0", "Moyenne": "FFFBEA", "Basse": "F0FDF4"}
            if priority in color_map:
                set_cell_bg(row[3], color_map[priority])

            for cell in row:
                for run in cell.paragraphs[0].runs:
                    run.font.size = Pt(9)

            # Alternance de couleur
            if i % 2 == 1:
                for j in [0, 1, 2, 4]:
                    set_cell_bg(row[j], "F7F7F5")

    # ── Sauvegarde ────────────────────────────────────────────────────────────
    filename = f"CDC_{client_name.replace(' ', '_')}_{session_id[:6]}.docx"
    filepath = os.path.join(OUTPUT_DIR, filename)
    doc.save(filepath)

    return filepath